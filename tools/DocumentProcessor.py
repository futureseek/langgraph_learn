#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
文档处理器 - 支持多种文档格式的解析和分块
支持Meta-Chunking策略的文档分块
"""

import os
import re
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional
import pypdf
import docx
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    MarkdownTextSplitter,
    PythonCodeTextSplitter
)
from langchain.docstore.document import Document
from openai import APIConnectionError, APITimeoutError
# 导入文本规范化器
from .text_normalizer import TextNormalizer
from .clean_think import clean_response
# 尝试导入可选依赖
try:
    import torch
    import torch.nn.functional as F
    from transformers import AutoModelForCausalLM, AutoTokenizer
except ImportError:
    torch = None
    F = None
    AutoModelForCausalLM = None
    AutoTokenizer = None
# 尝试导入 OpenAI（用于 API 调用模型）
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None
    
"""
默认的 API 配置（可直接在此处填入你的实际参数，或用环境变量覆盖）
优先级：函数入参 > 环境变量 > 下方默认常量
"""
DEFAULT_API_MODEL_NAME = os.getenv("OPENAI_MODEL_NAME") or "qwen3:1.7b"
DEFAULT_API_BASE_URL = os.getenv("OPENAI_BASE_URL") or "http://localhost:11434/v1"
DEFAULT_API_KEY = os.getenv("OPENAI_API_KEY") or "ollama"
DEFAULT_API_STREAMING = (os.getenv("OPENAI_STREAMING") or "True").lower() in {"1", "True", "yes"}

def load_qwen_model(model_name="qwen3:1.7b", max_retries=3):
    """
    加载Qwen3-1.7B模型和分词器
    
    Args:
        model_name: 模型名称
        max_retries: 最大重试次数
        
    Returns:
        tuple: (model, tokenizer) 或 (None, None) 如果加载失败
    """
    import time
    
    for attempt in range(max_retries):
        try:
            print(f"正在加载模型 {model_name} (尝试 {attempt + 1}/{max_retries})...")
            
            # 通过ModelScope加载模型
            tokenizer = AutoTokenizer.from_pretrained(
                model_name, 
                trust_remote_code=True,
                use_fast=False  # 使用慢速分词器以提高兼容性
            )
            
            model = AutoModelForCausalLM.from_pretrained(
                model_name,
                trust_remote_code=True,
                device_map="auto",  # 自动分配设备
                torch_dtype=torch.float16 if torch else None  # 使用float16以节省内存
            )
            
            # 设置为评估模式
            if model:
                model.eval()
                
            print(f"✅ 成功加载模型: {model_name}")
            return model, tokenizer
        except Exception as e:
            print(f"❌ 加载模型失败 (尝试 {attempt + 1}/{max_retries}): {e}")
            if attempt < max_retries - 1:
                print(f"等待5秒后重试...")
                time.sleep(5)
            else:
                print("已达到最大重试次数")
    
    print("💡 将使用简化版本的Meta-Chunking策略")
    return None, None

def create_api_client(api_key: Optional[str] = None,
                      base_url: Optional[str] = None,
                      model_name: Optional[str] = None):
    """
    创建用于通过 API 调用大模型的客户端（基于 openai SDK）。

    Args:
        api_key: API 密钥
        base_url: API 基础地址（如 ModelScope: https://api-inference.modelscope.cn/v1）
        model_name: 模型名称（服务端可解析的标识）

    Returns:
        tuple: (client, model_name) 或 (None, None) 如果创建失败或 openai 未安装
    """
    if OpenAI is None:
        print("警告: openai SDK 不可用，无法创建 API 客户端")
        return None, None
    try:
        # 使用优先级：入参 > 环境变量/默认常量
        resolved_api_key = api_key or DEFAULT_API_KEY
        resolved_base_url = base_url or DEFAULT_API_BASE_URL
        resolved_model_name = model_name or DEFAULT_API_MODEL_NAME

        client = OpenAI(api_key=resolved_api_key, base_url=resolved_base_url)
        return client, resolved_model_name
    except Exception as e:
        print(f"❌ 创建 API 客户端失败: {e}")
        return None, None

class MetaChunking:
    """Meta-Chunking策略实现类"""
    
    def __init__(self, model=None, tokenizer=None, api_client=None, api_model: Optional[str] = None, api_streaming: Optional[bool] = None):
        self.model = model
        self.tokenizer = tokenizer
        self.api_client = api_client
        self.api_model = api_model
        # 若未显式传入，则使用默认配置
        self.api_streaming = DEFAULT_API_STREAMING if api_streaming is None else api_streaming
        
    def get_prob_subtract(self, sentence1, sentence2, language):
        """计算两个句子的概率差"""
        try:
            # 优先使用 API 调用做决策
            if self.api_client and self.api_model:
                print("[MetaChunking][API] 使用 API 做相邻句对决策")
                if language == 'zh':
                    query = '''这是一个文本分块任务.你是一位文本分析专家，请根据提供的句子的逻辑结构和语义内容，从下面两种方案中选择一种分块方式：
                    1. 将"{}"分割成"{}"与"{}"两部分；
                    2. 将"{}"不进行分割，保持原形式；
                    只回答数字1或2,不要提供其他解释。/no_think'''.format(sentence1 + sentence2, sentence1, sentence2, sentence1 + sentence2)
                    prompt = "You are a helpful assistant.\n\n{}\n\nAssistant:".format(query)
                else:
                    query = '''This is a text chunking task. You are a text analysis expert. Please choose one of the following two options based on the logical structure and semantic content of the provided sentence:
                    1. Split "{}" into "{}" and "{}" two parts;
                    2. Keep "{}" unsplit in its original form;
                    Answer with 1 or 2 only./no_think'''.format(sentence1 + ' ' + sentence2, sentence1, sentence2, sentence1 + ' ' + sentence2)
                    prompt = "You are a helpful assistant.\n\n{}\n\nAssistant:".format(query)

                try:
                    resp = self.api_client.chat.completions.create(
                        model=self.api_model,
                        messages=[
                            {"role": "system", "content": "You are a helpful assistant."},
                            {"role": "user", "content": prompt},
                        ],
                        temperature=0,
                        stream=True,
                        timeout=8,   # ⏱️ 设置超时时间 8 秒
                        max_tokens=10
                    )
                    if True:
                        # 流式拼接内容（兼容不同提供方字段）
                        parts = []
                        for ch in resp:
                            try:
                                delta = ch.choices[0].delta
                                seg = getattr(delta, 'content', None)
                                if seg is None and isinstance(delta, dict):
                                    seg = delta.get('content')
                                if seg:
                                    parts.append(seg)
                            except Exception:
                                # 兼容部分提供方使用 message.content
                                try:
                                    seg2 = ch.choices[0].message.content
                                    if seg2:
                                        parts.append(seg2)
                                except Exception:
                                    pass
                        content = ("".join(parts)).strip()
                        content = clean_response(content)
                    
                    # 规范化：只取第一个字符中的 1/2
                    decision = '1' if '1' in content[:3] and '2' not in content[:3] else (
                        '2' if '2' in content[:3] else content[:1]
                    )
                    if decision == '2':
                        # 选择不分割 -> 返回正值（> threshold 视为不分割）
                        print(f"[MetaChunking][API] 回复: {content!r} -> 决策=2(不分割) | 返回分数=1")
                        return 1
                    elif decision == '1':
                        # 选择分割 -> 返回负值（<= threshold 视为分割）
                        print(f"[MetaChunking][API] 回复: {content!r} -> 决策=1(分割) | 返回分数=-1")
                        return -1
                    else:
                        # 无法解析，返回中性
                        print(f"[MetaChunking][API] 无法解析回复: {content!r} | 返回分数=0")
                        return 0
                except (APITimeoutError, APIConnectionError) as e:
                    print(f"❌ API 请求超时/连接错误: {e}")
                    return 0  # 回退，不影响主流程
                except Exception as e:
                    print(f"API 决策失败，回退本地/简化逻辑: {e}")
                    # 回退到本地逻辑

            # 检查是否有可用的本地模型和分词器
            if not self.model or not self.tokenizer:
                print("[MetaChunking][Local] 本地模型或分词器不可用，返回分数0")
                return 0  # 返回默认值
            
            if language == 'zh':
                query = '''这是一个文本分块任务.你是一位文本分析专家，请根据提供的句子的逻辑结构和语义内容，从下面两种方案中选择一种分块方式：
                1. 将"{}"分割成"{}"与"{}"两部分；
                2. 将"{}"不进行分割，保持原形式；
                请回答1或2。'''.format(sentence1 + sentence2, sentence1, sentence2, sentence1 + sentence2)
                prompt = "You are a helpful assistant.\n\n{}\n\nAssistant:".format(query)
            else:
                query = '''This is a text chunking task. You are a text analysis expert. Please choose one of the following two options based on the logical structure and semantic content of the provided sentence:
                1. Split "{}" into "{}" and "{}" two parts;
                2. Keep "{}" unsplit in its original form;
                Please answer 1 or 2.'''.format(sentence1 + ' ' + sentence2, sentence1, sentence2, sentence1 + ' ' + sentence2)
                prompt = "You are a helpful assistant.\n\n{}\n\nAssistant:".format(query)
            
            prompt_ids = self.tokenizer.encode(prompt, return_tensors='pt').to(self.model.device)
            input_ids = prompt_ids
            output_ids = self.tokenizer.encode(['1','2'], return_tensors='pt').to(self.model.device)
            
            with torch.no_grad():
                outputs = self.model(input_ids)
                next_token_logits = outputs.logits[:, -1, :]
                token_probs = F.softmax(next_token_logits, dim=-1)
            
            next_token_id_0 = output_ids[:, 0].unsqueeze(0)
            next_token_prob_0 = token_probs[:, next_token_id_0].item()      
            next_token_id_1 = output_ids[:, 1].unsqueeze(0)
            next_token_prob_1 = token_probs[:, next_token_id_1].item()  
            prob_subtract = next_token_prob_1 - next_token_prob_0
            
            return prob_subtract
        except Exception as e:
            print(f"计算概率差失败: {e}")
            return 0  # 返回默认值
    
    def _fallback_chunking(self, text):
        """回退到递归字符分割"""
        try:
            # 使用简单的字符分割作为回退方案
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                separators=["\n\n", "\n", ".", "!", "?", " ", ""]
            )
            return text_splitter.split_text(text)
        except Exception as e:
            print(f"回退分块也失败了，返回原始文本: {e}")
            return [text]
    
    def split_text_by_punctuation(self, text, language): 
        """根据标点符号分割文本"""
        try:
            # 检查jieba是否可用
            import jieba
            jieba_available = True
        except ImportError:
            jieba_available = False
            print("警告: jieba模块不可用，使用简单分割")
        
        if language == 'zh': 
            # 更健壮的中文断句：支持更多标点与换行
            puncts = {"。", "！", "？", "；", "：", "，", "．"}
            if jieba_available:
                words = list(jieba.cut(text, cut_all=False))
                sentences = []
                buf = ""
                for w in words:
                    if w in puncts or w == "\n":
                        if buf.strip():
                            sentences.append(buf.strip() + (w if w != "\n" else ""))
                        buf = ""
                    else:
                        buf += w
                if buf.strip():
                    sentences.append(buf.strip())
                # 如果仍然过少，使用正则进一步切
                if len(sentences) <= 1:
                    import re
                    parts = re.split(r'[。！？；：，．\n]+', text)
                    sentences = [s.strip() for s in parts if s.strip()]
                return sentences
            else:
                # 纯正则断句
                import re
                parts = re.split(r'[。！？；：，．\n]+', text)
                return [s.strip() for s in parts if s.strip()]
        else:
            try:
                from nltk.tokenize import sent_tokenize
                full_segments = sent_tokenize(text)
            except ImportError:
                # 如果nltk不可用，使用简单分割
                import re
                full_segments = re.split(r'[.!?]+', text)
            
            ret = []
            for item in full_segments:
                item_l = item.strip().split(' ')
                if len(item_l) > 512:
                    if len(item_l) > 1024:
                        item = ' '.join(item_l[:256]) + "..."
                    else:
                        item = ' '.join(item_l[:512]) + "..."
                ret.append(item)
            return ret
    
    def perplexity_chunking(self, text, threshold: float = 0.0, language: str = 'en', target_length: Optional[int] = None) -> List[str]:
        """基于困惑度（PPL）的分块：
        - 句级切分，计算 PPL 序列 PPL(x1..xn) 相对前文上下文
        - 识别局部极小值作为候选边界；根据阈值与目标长度进行边界筛选
        - 长文本采用滑动窗口近似 KV-cache：仅保留最近窗口上下文
        """
        try:
            segments = self.split_text_by_punctuation(text, language)
            segments = [s for s in segments if s.strip()]
            if len(segments) <= 1:
                print(f"[Chunk][PPL] 段数={len(segments)}，直接返回整体")
                return [text]

            def _calc_ppl_seq_with_transformers(seg_list: List[str]) -> List[float]:
                if not (self.model and self.tokenizer and torch is not None):
                    return []
                ppl_vals: List[float] = []
                context = ""
                # 简化滑动窗口（字符级，近似 KV-cache）：
                max_ctx_chars = 3000
                for i, sent in enumerate(seg_list):
                    if context:
                        prompt = context + (" " if language != 'zh' else "") + sent
                    else:
                        prompt = sent
                    input_ids = self.tokenizer(prompt, return_tensors='pt').input_ids.to(self.model.device)
                    with torch.no_grad():
                        out = self.model(input_ids)
                        logits = out.logits[:, :-1, :]
                        labels = input_ids[:, 1:]
                        nll = torch.nn.functional.cross_entropy(logits.reshape(-1, logits.size(-1)), labels.reshape(-1), reduction='mean')
                    ppl = float(torch.exp(nll).item()) if hasattr(torch, 'exp') else float(nll.item())
                    ppl_vals.append(ppl)
                    context = (context + (" " if language != 'zh' else "") + sent) if context else sent
                    if len(context) > max_ctx_chars:
                        context = context[-max_ctx_chars:]
                return ppl_vals

            def _calc_ppl_seq_proxy(seg_list: List[str]) -> List[float]:
                # 无模型时的 proxy：用逆长度和标点密度混合，越难预测→越高 PPL
                vals = []
                for s in seg_list:
                    ln = max(1, len(s))
                    punct = sum(c in '，。！？；:：,.!?;…' for c in s) + 1
                    vals.append((1000.0/ln) + (1.0/punct))
                return vals

            ppl_seq = _calc_ppl_seq_with_transformers(segments)
            if not ppl_seq:
                print("[Chunk][PPL] 无可用模型，使用 proxy PPL")
                ppl_seq = _calc_ppl_seq_proxy(segments)

            # 寻找局部极小值作为边界候选
            minima = set()
            for i in range(1, len(ppl_seq)-1):
                if ppl_seq[i] <= ppl_seq[i-1] and ppl_seq[i] <= ppl_seq[i+1]:
                    minima.add(i)

            # 根据阈值筛选（相对全局均值/方差）
            import statistics
            mean = statistics.mean(ppl_seq)
            stdev = statistics.pstdev(ppl_seq) if len(ppl_seq) > 1 else 0.0
            cut = mean - threshold * (stdev if stdev > 0 else 1.0)

            boundaries = []
            for i in sorted(minima):
                if ppl_seq[i] <= cut:
                    boundaries.append(i)

            # 基于边界生成元块
            chunks: List[str] = []
            last = 0
            for b in boundaries:
                piece = segments[last:b+1]
                chunk = piece[0]
                for s in piece[1:]:
                    chunk = (chunk + (" " if language != 'zh' else "") + s)
                chunks.append(chunk)
                last = b+1
            if last < len(segments):
                rest = segments[last:]
                chunk = rest[0]
                for s in rest[1:]:
                    chunk = (chunk + (" " if language != 'zh' else "") + s)
                chunks.append(chunk)

            if not chunks:
                return [text]

            # 动态合并使块长接近目标长度
            if target_length and len(chunks) > 1:
                merged: List[str] = []
                acc = ""
                for ck in chunks:
                    if not acc:
                        acc = ck
                        continue
                    if len(acc) + len(ck) <= target_length:
                        acc = acc + ((" " if language != 'zh' else "") + ck)
                    else:
                        merged.append(acc)
                        acc = ck
                if acc:
                    merged.append(acc)
                print(f"[Chunk][PPL] 依据目标长度={target_length} 合并后块数={len(merged)}")
                return merged

            return chunks
        except Exception as e:
            print(f"困惑度分块失败，回退到默认分块: {e}")
            return self._fallback_chunking(text)
    
    def prob_subtract_chunking(self, text, threshold=0, language='en', target_length: Optional[int] = None) -> List[str]:
        """基于边缘采样的分块（Margin Sampling）：
        - 元块 Xj 为当前累积块，xi 为下一句
        - 使用 API 决策是否将 xi 合并进 Xj（score 与动态阈值对比）
        - 可选 target_length：在合并时与事后动态合并，尽量靠近目标块长
        """
        try:
            segments = self.split_text_by_punctuation(text, language)
            segments = [item for item in segments if item.strip()]

            if len(segments) <= 1:
                print(f"[Chunk][ProbSubtract] 段数={len(segments)}，直接返回整体")
                return [text]

            # 若既无 API 也无本地模型，则回退到简化的长度逻辑
            if not ((self.api_client and self.api_model) or (self.model and self.tokenizer)):
                print(f"[Chunk][ProbSubtract] 无 API/本地模型，使用长度回退逻辑 | 段数={len(segments)}")
                chunks = []
                current_chunk = ""
                for segment in segments:
                    if current_chunk and len(current_chunk) + len(segment) > 800:
                        print(f"[Chunk][LenFallback] 触发切分 | 当前块长度={len(current_chunk)} | 新段长度={len(segment)}")
                        chunks.append(current_chunk)
                        current_chunk = segment
                    else:
                        if language == 'zh':
                            current_chunk += segment
                        else:
                            current_chunk += (" " if current_chunk else "") + segment
                if current_chunk:
                    chunks.append(current_chunk)
                print(f"[Chunk][LenFallback] 完成 | 块数={len(chunks)}")
                return chunks if chunks else [text]

            # 使用相邻句对 + 动态阈值（最近5个的均值）
            print(f"[Chunk][ProbSubtract] 使用相邻句对 + 动态阈值 | 段数={len(segments)} | 决策来源={'API' if (self.api_client and self.api_model) else 'Local'}")
            chunks: List[str] = []
            tmp_chunk = ""
            scores: List[float] = []
            dynamic_threshold: float = 0.0

            for i, sentence in enumerate(segments):
                if tmp_chunk == "":
                    tmp_chunk = sentence
                    continue

                # 使用元块 Xj(=tmp_chunk) vs 当前句 xi(=sentence)
                score = self.get_prob_subtract(tmp_chunk, sentence, language)
                scores.append(score)

                # 更新动态阈值（最近5个的平均值）
                recent = scores[-5:]
                dynamic_threshold = sum(recent) / len(recent)

                # 若提供目标块长，优先防止块无限增大
                force_split = False
                if target_length and len(tmp_chunk) >= max(1, int(1.2 * target_length)):
                    force_split = True

                if not force_split and score > dynamic_threshold:
                    # 合并到当前块 Xj
                    if language == 'zh':
                        tmp_chunk += sentence
                    else:
                        tmp_chunk += " " + sentence
                    print(f"[Chunk][Step i={i}] score={score:.4f} | thr={dynamic_threshold:.4f} -> 合并")
                else:
                    # 新起一个块
                    chunks.append(tmp_chunk)
                    print(f"[Chunk][Step i={i}] score={score:.4f} | thr={dynamic_threshold:.4f} | 执行动作=切分 | 当前块数={len(chunks)}")
                    tmp_chunk = sentence

            if tmp_chunk:
                chunks.append(tmp_chunk)
            print(f"[Chunk][ProbSubtract] 完成 | 最终块数={len(chunks)}")

            # 事后动态合并：尽量靠近 target_length
            if target_length and len(chunks) > 1:
                merged: List[str] = []
                acc = ""
                for ck in chunks:
                    if not acc:
                        acc = ck
                        continue
                    if len(acc) + len(ck) <= target_length:
                        if language == 'zh':
                            acc += ck
                        else:
                            acc += " " + ck
                    else:
                        merged.append(acc)
                        acc = ck
                if acc:
                    merged.append(acc)
                print(f"[Chunk][DynamicMerge] 依据目标长度={target_length} 合并后块数={len(merged)}")
                return merged if merged else chunks

            return chunks if chunks else [text]
        except Exception as e:
            print(f"概率差分块失败，回退到默认分块: {e}")
            return self._fallback_chunking(text)
    
    def semantic_chunking(self, text, breakpoint_percentile_threshold=73, language='en'):
        """基于语义的分块策略（需要llama_index库）"""
        try:
            # 这里我们实现一个简化的语义分块逻辑
            # 实际的语义分块需要使用嵌入模型来识别语义边界
            # 由于依赖问题，我们使用简化版本
            segments = self.split_text_by_punctuation(text, language)
            segments = [item for item in segments if item.strip()]
            
            if len(segments) <= 1:
                return [text]
            
            # 简化的语义分块：根据句子长度和标点符号进行分组
            chunks = []
            current_chunk = ""
            
            for segment in segments:
                if len(current_chunk) + len(segment) > 1000:  # 假设最大块大小为1000字符
                    if current_chunk:
                        chunks.append(current_chunk)
                    current_chunk = segment
                else:
                    if language == 'zh':
                        current_chunk += segment
                    else:
                        current_chunk += " " + segment if current_chunk else segment
            
            if current_chunk:
                chunks.append(current_chunk)
            
            return chunks if chunks else [text]
        except Exception as e:
            print(f"语义分块失败，回退到默认分块: {e}")
            return self._fallback_chunking(text)

class DocumentProcessor:
    """
    文档处理器 - 支持多种文档格式的解析和分块
    """
    
    def __init__(self, 
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200,
                 separators: Optional[List[str]] = None,
                 meta_chunking_strategy: Optional[str] = None,
                 meta_model = None,
                 meta_tokenizer = None,
                 api_client = None,
                 api_model: Optional[str] = None):
        """
        初始化文档处理器
        
        Args:
            chunk_size: 文档块大小
            chunk_overlap: 块之间的重叠
            separators: 自定义分隔符
            meta_chunking_strategy: Meta-Chunking策略 ("perplexity", "prob_subtract", "semantic", None)
            meta_model: 用于Meta-Chunking的模型
            meta_tokenizer: 用于Meta-Chunking的分词器
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # 支持的文件格式
        self.supported_formats = {
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.py': self._process_python,
            '.js': self._process_javascript,
            '.java': self._process_java,
            '.cpp': self._process_cpp,
            '.c': self._process_cpp,
            '.h': self._process_cpp,
            '.hpp': self._process_cpp,
            '.html': self._process_html,
            '.xml': self._process_xml,
            '.json': self._process_json,
            '.csv': self._process_csv,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.log': self._process_text,
            '.cfg': self._process_text,
            '.ini': self._process_text,
            '.yaml': self._process_text,
            '.yml': self._process_text,
        }
        
        # 初始化不同类型的文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators or ["\n\n", "\n", " ", ""]
        )
        
        self.markdown_splitter = MarkdownTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        
        self.code_splitter = PythonCodeTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        
        # 初始化Meta-Chunking
        self.meta_chunking_strategy = meta_chunking_strategy
        if meta_chunking_strategy:
            self.meta_chunking = MetaChunking(meta_model, meta_tokenizer, api_client=api_client, api_model=api_model)
        else:
            self.meta_chunking = None
        
        # 初始化文本规范化器
        self.text_normalizer = TextNormalizer()
        
    def process_file(self, file_path: str) -> List[Document]:
        """
        处理单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            List[Document]: 处理后的文档块列表
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            # 获取文件扩展名
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext not in self.supported_formats:
                raise ValueError(f"不支持的文件格式: {file_ext}")
            
            # 获取文件基本信息
            file_stats = os.stat(file_path)
            base_metadata = {
                "source": os.path.abspath(file_path),
                "file_name": os.path.basename(file_path),
                "file_type": file_ext,
                "file_size": file_stats.st_size,
                "modified_time": file_stats.st_mtime
            }
            
            # 调用对应的处理函数
            processor_func = self.supported_formats[file_ext]
            documents = processor_func(file_path, base_metadata)
            
            print(f"✅ 成功处理文件: {os.path.basename(file_path)} - {len(documents)} 个块")
            return documents
            
        except Exception as e:
            print(f"❌ 处理文件失败 {file_path}: {e}")
            return []
    
    def process_directory(self, directory_path: str, 
                         recursive: bool = True,
                         file_pattern: Optional[str] = None) -> List[Document]:
        """
        处理目录下的所有支持文件
        
        Args:
            directory_path: 目录路径
            recursive: 是否递归处理子目录
            file_pattern: 文件名模式过滤（正则表达式）
            
        Returns:
            List[Document]: 所有文档块列表
        """
        all_documents = []
        
        try:
            directory = Path(directory_path)
            if not directory.exists() or not directory.is_dir():
                raise ValueError(f"目录不存在或不是目录: {directory_path}")
            
            # 获取文件列表
            if recursive:
                files = directory.rglob("*")
            else:
                files = directory.glob("*")
            
            processed_files = 0
            
            for file_path in files:
                if not file_path.is_file():
                    continue
                
                # 检查文件扩展名
                if file_path.suffix.lower() not in self.supported_formats:
                    continue
                
                # 应用文件名模式过滤
                if file_pattern and not re.match(file_pattern, file_path.name):
                    continue
                
                # 处理文件
                documents = self.process_file(str(file_path))
                all_documents.extend(documents)
                processed_files += 1
            
            print(f"✅ 目录处理完成: {processed_files} 个文件, 共 {len(all_documents)} 个文档块")
            return all_documents
            
        except Exception as e:
            print(f"❌ 处理目录失败 {directory_path}: {e}")
            return []
    
    def process_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """
        直接处理文本内容
        
        Args:
            text: 文本内容
            metadata: 元数据
            
        Returns:
            List[Document]: 文档块列表
        """
        try:
            if not text.strip():
                return []
            
            # 文本规范化
            normalized_text = self.text_normalizer.normalize_text(text)
            
            # 使用Meta-Chunking或默认分割器
            if self.meta_chunking and self.meta_chunking_strategy:
                if self.meta_chunking_strategy == 'perplexity':
                    chunks = self.meta_chunking.perplexity_chunking(normalized_text)
                elif self.meta_chunking_strategy == 'prob_subtract':
                    chunks = self.meta_chunking.prob_subtract_chunking(normalized_text)
                elif self.meta_chunking_strategy == 'semantic':
                    chunks = self.meta_chunking.semantic_chunking(normalized_text)
                else:
                    chunks = self.text_splitter.split_text(normalized_text)
            else:
                chunks = self.text_splitter.split_text(normalized_text)
            
            # 创建文档对象
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = metadata.copy() if metadata else {}
                doc_metadata.update({
                    "chunk_index": i,
                    "chunk_size": len(chunk)
                })
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            
            return documents
            
        except Exception as e:
            print(f"❌ 处理文本失败: {e}")
            return []
    
    def _process_text(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理纯文本文件"""
        encodings = ['utf-8', 'gbk', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError(f"无法解码文件: {file_path}")
        
        # 文本规范化
        normalized_content = self.text_normalizer.normalize_text(content)
        
        # 使用Meta-Chunking或默认分割器
        if self.meta_chunking and self.meta_chunking_strategy:
            if self.meta_chunking_strategy == 'perplexity':
                chunks = self.meta_chunking.perplexity_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'prob_subtract':
                chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'semantic':
                chunks = self.meta_chunking.semantic_chunking(normalized_content)
            else:
                chunks = self.text_splitter.split_text(normalized_content)
        else:
            chunks = self.text_splitter.split_text(normalized_content)
        
        return self._create_documents(chunks, base_metadata)
    
    def _process_markdown(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 文本规范化
        normalized_content = self.text_normalizer.normalize_text(content)
        
        # 使用Meta-Chunking或默认分割器
        if self.meta_chunking and self.meta_chunking_strategy:
            if self.meta_chunking_strategy == 'perplexity':
                chunks = self.meta_chunking.perplexity_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'prob_subtract':
                chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'semantic':
                chunks = self.meta_chunking.semantic_chunking(normalized_content)
            else:
                chunks = self.markdown_splitter.split_text(normalized_content)
        else:
            chunks = self.markdown_splitter.split_text(normalized_content)
        
        return self._create_documents(chunks, base_metadata)
    
    def _process_python(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理Python代码文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 文本规范化
        normalized_content = self.text_normalizer.normalize_text(content)
        
        # 使用Meta-Chunking或默认分割器
        if self.meta_chunking and self.meta_chunking_strategy:
            if self.meta_chunking_strategy == 'perplexity':
                chunks = self.meta_chunking.perplexity_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'prob_subtract':
                chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'semantic':
                chunks = self.meta_chunking.semantic_chunking(normalized_content)
            else:
                chunks = self.code_splitter.split_text(normalized_content)
        else:
            chunks = self.code_splitter.split_text(normalized_content)
        
        return self._create_documents(chunks, base_metadata)
    
    def _process_javascript(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理JavaScript文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 文本规范化
        normalized_content = self.text_normalizer.normalize_text(content)
        
        # 使用Meta-Chunking或默认分割器
        if self.meta_chunking and self.meta_chunking_strategy:
            if self.meta_chunking_strategy == 'perplexity':
                chunks = self.meta_chunking.perplexity_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'prob_subtract':
                chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'semantic':
                chunks = self.meta_chunking.semantic_chunking(normalized_content)
            else:
                chunks = self.text_splitter.split_text(normalized_content)
        else:
            chunks = self.text_splitter.split_text(normalized_content)
        
        return self._create_documents(chunks, base_metadata)
    
    def _process_java(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理Java文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 文本规范化
        normalized_content = self.text_normalizer.normalize_text(content)
        
        # 使用Meta-Chunking或默认分割器
        if self.meta_chunking and self.meta_chunking_strategy:
            if self.meta_chunking_strategy == 'perplexity':
                chunks = self.meta_chunking.perplexity_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'prob_subtract':
                chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'semantic':
                chunks = self.meta_chunking.semantic_chunking(normalized_content)
            else:
                chunks = self.text_splitter.split_text(normalized_content)
        else:
            chunks = self.text_splitter.split_text(normalized_content)
        
        return self._create_documents(chunks, base_metadata)
    
    def _process_cpp(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理C/C++文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 文本规范化
        normalized_content = self.text_normalizer.normalize_text(content)
        
        # 使用Meta-Chunking或默认分割器
        if self.meta_chunking and self.meta_chunking_strategy:
            if self.meta_chunking_strategy == 'perplexity':
                chunks = self.meta_chunking.perplexity_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'prob_subtract':
                chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'semantic':
                chunks = self.meta_chunking.semantic_chunking(normalized_content)
            else:
                chunks = self.text_splitter.split_text(normalized_content)
        else:
            chunks = self.text_splitter.split_text(normalized_content)
        
        return self._create_documents(chunks, base_metadata)
    
    def _process_html(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理HTML文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 简单去除HTML标签（可以后续使用BeautifulSoup改进）
        import re
        content = re.sub(r'<[^>]+>', '', content)
        content = re.sub(r'\s+', ' ', content).strip()
        
        # 文本规范化
        normalized_content = self.text_normalizer.normalize_text(content)
        
        # 使用Meta-Chunking或默认分割器
        if self.meta_chunking and self.meta_chunking_strategy:
            if self.meta_chunking_strategy == 'perplexity':
                chunks = self.meta_chunking.perplexity_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'prob_subtract':
                chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'semantic':
                chunks = self.meta_chunking.semantic_chunking(normalized_content)
            else:
                chunks = self.text_splitter.split_text(normalized_content)
        else:
            chunks = self.text_splitter.split_text(normalized_content)
        
        return self._create_documents(chunks, base_metadata)
    
    def _process_xml(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理XML文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 文本规范化
        normalized_content = self.text_normalizer.normalize_text(content)
        
        # 使用Meta-Chunking或默认分割器
        if self.meta_chunking and self.meta_chunking_strategy:
            if self.meta_chunking_strategy == 'perplexity':
                chunks = self.meta_chunking.perplexity_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'prob_subtract':
                chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'semantic':
                chunks = self.meta_chunking.semantic_chunking(normalized_content)
            else:
                chunks = self.text_splitter.split_text(normalized_content)
        else:
            chunks = self.text_splitter.split_text(normalized_content)
        
        return self._create_documents(chunks, base_metadata)
    
    def _process_json(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理JSON文件"""
        import json
        
        with open(file_path, 'r', encoding='utf-8') as f:
            try:
                data = json.load(f)
                content = json.dumps(data, ensure_ascii=False, indent=2)
            except json.JSONDecodeError:
                # 如果JSON解析失败，作为普通文本处理
                f.seek(0)
                content = f.read()
        
        # 文本规范化
        normalized_content = self.text_normalizer.normalize_text(content)
        
        # 使用Meta-Chunking或默认分割器
        if self.meta_chunking and self.meta_chunking_strategy:
            if self.meta_chunking_strategy == 'perplexity':
                chunks = self.meta_chunking.perplexity_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'prob_subtract':
                chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'semantic':
                chunks = self.meta_chunking.semantic_chunking(normalized_content)
            else:
                chunks = self.text_splitter.split_text(normalized_content)
        else:
            chunks = self.text_splitter.split_text(normalized_content)
        
        return self._create_documents(chunks, base_metadata)
    
    def _process_csv(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理CSV文件"""
        import csv
        
        content_lines = []
        
        with open(file_path, 'r', encoding='utf-8') as f:
            try:
                reader = csv.reader(f)
                for row in reader:
                    content_lines.append(",".join(row))
            except Exception:
                # 如果CSV解析失败，作为普通文本处理
                f.seek(0)
                content_lines = f.readlines()
        
        content = "\n".join(content_lines)
        
        # 文本规范化
        normalized_content = self.text_normalizer.normalize_text(content)
        
        # 使用Meta-Chunking或默认分割器
        if self.meta_chunking and self.meta_chunking_strategy:
            if self.meta_chunking_strategy == 'perplexity':
                chunks = self.meta_chunking.perplexity_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'prob_subtract':
                chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
            elif self.meta_chunking_strategy == 'semantic':
                chunks = self.meta_chunking.semantic_chunking(normalized_content)
            else:
                chunks = self.text_splitter.split_text(normalized_content)
        else:
            chunks = self.text_splitter.split_text(normalized_content)
        
        return self._create_documents(chunks, base_metadata)
    
    def _process_pdf(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理PDF文件"""
        try:
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                content = ""
                
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        content += f"\n--- 第 {page_num + 1} 页 ---\n"
                        content += page_text + "\n"
            
            if not content.strip():
                raise ValueError("PDF文件无法提取文本内容")
            
            # 文本规范化
            normalized_content = self.text_normalizer.normalize_text(content)
            
            # 使用Meta-Chunking或默认分割器
            if self.meta_chunking and self.meta_chunking_strategy:
                if self.meta_chunking_strategy == 'perplexity':
                    chunks = self.meta_chunking.perplexity_chunking(normalized_content)
                elif self.meta_chunking_strategy == 'prob_subtract':
                    chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
                elif self.meta_chunking_strategy == 'semantic':
                    chunks = self.meta_chunking.semantic_chunking(normalized_content)
                else:
                    chunks = self.text_splitter.split_text(normalized_content)
            else:
                chunks = self.text_splitter.split_text(normalized_content)
            
            return self._create_documents(chunks, base_metadata)
            
        except Exception as e:
            print(f"❌ PDF处理失败: {e}")
            return []
    
    def _process_docx(self, file_path: str, base_metadata: Dict) -> List[Document]:
        """处理Word文档"""
        try:
            doc = docx.Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        content += row_text + "\n"
            
            if not content.strip():
                raise ValueError("Word文档无法提取文本内容")
            
            # 文本规范化
            normalized_content = self.text_normalizer.normalize_text(content)
            
            # 使用Meta-Chunking或默认分割器
            if self.meta_chunking and self.meta_chunking_strategy:
                if self.meta_chunking_strategy == 'perplexity':
                    chunks = self.meta_chunking.perplexity_chunking(normalized_content)
                elif self.meta_chunking_strategy == 'prob_subtract':
                    chunks = self.meta_chunking.prob_subtract_chunking(normalized_content)
                elif self.meta_chunking_strategy == 'semantic':
                    chunks = self.meta_chunking.semantic_chunking(normalized_content)
                else:
                    chunks = self.text_splitter.split_text(normalized_content)
            else:
                chunks = self.text_splitter.split_text(normalized_content)
            
            return self._create_documents(chunks, base_metadata)
            
        except Exception as e:
            print(f"❌ Word文档处理失败: {e}")
            return []
    
    def _create_documents(self, chunks: List[str], base_metadata: Dict) -> List[Document]:
        """创建文档对象列表"""
        documents = []
        
        for i, chunk in enumerate(chunks):
            if not chunk.strip():
                continue
            
            doc_metadata = base_metadata.copy()
            doc_metadata.update({
                "chunk_index": i,
                "chunk_size": len(chunk),
                "total_chunks": len(chunks)
            })
            
            documents.append(Document(
                page_content=chunk,
                metadata=doc_metadata
            ))
        
        return documents
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式列表"""
        return list(self.supported_formats.keys())
    
    def estimate_chunks(self, text_length: int) -> int:
        """估算文本块数量"""
        return max(1, text_length // (self.chunk_size - self.chunk_overlap))


def create_document_processor(chunk_size: int = 1000,
                            chunk_overlap: int = 200,
                            meta_chunking_strategy: Optional[str] = None,
                            meta_model = None,
                            meta_tokenizer = None,
                            api_client = None,
                            api_model: Optional[str] = None) -> DocumentProcessor:
    """
    创建文档处理器实例
    
    Args:
        chunk_size: 文档块大小
        chunk_overlap: 块重叠大小
        meta_chunking_strategy: Meta-Chunking策略
        meta_model: 用于Meta-Chunking的模型
        meta_tokenizer: 用于Meta-Chunking的分词器
        
    Returns:
        DocumentProcessor: 文档处理器实例
    """
    return DocumentProcessor(
        chunk_size=chunk_size, 
        chunk_overlap=chunk_overlap,
        meta_chunking_strategy=meta_chunking_strategy,
        meta_model=meta_model,
        meta_tokenizer=meta_tokenizer,
        api_client=api_client,
        api_model=api_model
    )


# 测试主函数
def main():
    # 测试文本
    test_text = (
        "这是第一句中文测试文本。"
        "这是第二句，用于测试分块效果！"
        "第三句继续测试？"
        "第四句用于检验Meta-Chunking是否正常工作。"
    )

    print("===== 1. 测试加载本地模型 =====")
    model, tokenizer = load_qwen_model()
    if model and tokenizer:
        print("✅ 本地模型可用")
    else:
        print("⚠️ 本地模型不可用，将使用API或回退逻辑")

    print("\n===== 2. 测试创建API客户端 =====")
    api_client, api_model_name = create_api_client()
    if api_client:
        print(f"✅ API客户端可用，模型: {api_model_name}")
    else:
        print("⚠️ API客户端不可用")

    print("\n===== 3. 测试MetaChunking分块 =====")
    # 初始化MetaChunking
    meta_chunker = MetaChunking(
        model=model,
        tokenizer=tokenizer,
        api_client=api_client,
        api_model=api_model_name
    )

    # 测试prob_subtract_chunking
    chunks = meta_chunker.prob_subtract_chunking(test_text, language='zh', target_length=50)
    print(f"分块结果 ({len(chunks)} 块):")
    for i, c in enumerate(chunks, 1):
        print(f"[块{i}] {c}")

    print("\n===== 4. 测试DocumentProcessor处理文本 =====")
    doc_processor = DocumentProcessor(
        chunk_size=50,
        chunk_overlap=10,
        meta_chunking_strategy='prob_subtract',
        meta_model=model,
        meta_tokenizer=tokenizer,
        api_client=api_client,
        api_model=api_model_name
    )
    documents = doc_processor.process_text(test_text)
    print(f"处理后的Document对象数量: {len(documents)}")
    for i, doc in enumerate(documents, 1):
        print(f"[Document {i}] {doc.page_content}")

if __name__ == "__main__":
    main()
    
